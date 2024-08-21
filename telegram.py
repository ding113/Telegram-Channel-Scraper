import asyncio
import aiohttp
from bs4 import BeautifulSoup
import re
import time
import html
from datetime import datetime
import logging
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
import json
import yaml
import argparse
from abc import ABC, abstractmethod
import csv
import markdown
from fpdf import FPDF
import xlsxwriter
from docx import Document
from tqdm import tqdm
import os
import sys

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("telegram_scraper.log", encoding="utf-8"),
    ],
)

logger = logging.getLogger(__name__)

@dataclass
class TelegramPost:
    message_id: str
    date: str
    text: str
    photo_url: str

class OutputProcessor(ABC):
    @abstractmethod
    def save(self, posts: List[TelegramPost], filename: str, delimiter: str = None):
        pass

class JSONOutputProcessor(OutputProcessor):
    def save(self, posts: List[TelegramPost], filename: str, delimiter: str = None):
        with open(filename, "w", encoding="utf-8") as jsonfile:
            json.dump([post.__dict__ for post in posts], jsonfile, ensure_ascii=False, indent=2)

class TXTOutputProcessor(OutputProcessor):
    def save(self, posts: List[TelegramPost], filename: str, delimiter: str = "\n"):
        with open(filename, "w", encoding="utf-8") as txtfile:
            for post in posts:
                txtfile.write(f"ID: {post.message_id}{delimiter}")
                txtfile.write(f"Date: {post.date}{delimiter}")
                txtfile.write(f"Text: {post.text}{delimiter}")
                if post.photo_url:
                    txtfile.write(f"Photo URL: {post.photo_url}{delimiter}")
                txtfile.write(f"{delimiter}")

class MarkdownOutputProcessor(OutputProcessor):
    def save(self, posts: List[TelegramPost], filename: str, delimiter: str = None):
        with open(filename, "w", encoding="utf-8") as mdfile:
            for post in posts:
                mdfile.write(f"## Message ID: {post.message_id}\n\n")
                mdfile.write(f"**Date:** {post.date}\n\n")
                mdfile.write(f"**Text:**\n\n{post.text}\n\n")
                if post.photo_url:
                    mdfile.write(f"**Photo:** ![Photo]({post.photo_url})\n\n")
                mdfile.write("---\n\n")

class PDFOutputProcessor(OutputProcessor):
    def save(self, posts: List[TelegramPost], filename: str, delimiter: str = None):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for post in posts:
            pdf.cell(200, 10, txt=f"Message ID: {post.message_id}", ln=1)
            pdf.cell(200, 10, txt=f"Date: {post.date}", ln=1)
            pdf.multi_cell(0, 10, txt=f"Text: {post.text}")
            if post.photo_url:
                pdf.cell(200, 10, txt=f"Photo URL: {post.photo_url}", ln=1)
            pdf.ln(10)
        pdf.output(filename)

class HTMLOutputProcessor(OutputProcessor):
    def save(self, posts: List[TelegramPost], filename: str, delimiter: str = None):
        with open(filename, "w", encoding="utf-8") as htmlfile:
            htmlfile.write("<html><body>")
            for post in posts:
                htmlfile.write(f"<h2>Message ID: {post.message_id}</h2>")
                htmlfile.write(f"<p><strong>Date:</strong> {post.date}</p>")
                htmlfile.write(f"<p><strong>Text:</strong><br>{post.text}</p>")
                if post.photo_url:
                    htmlfile.write(f"<p><strong>Photo:</strong><br><img src='{post.photo_url}' alt='Photo'></p>")
                htmlfile.write("<hr>")
            htmlfile.write("</body></html>")

class XLSXOutputProcessor(OutputProcessor):
    def save(self, posts: List[TelegramPost], filename: str, delimiter: str = None):
        workbook = xlsxwriter.Workbook(filename)
        worksheet = workbook.add_worksheet()
        headers = ["Message ID", "Date", "Text", "Photo URL"]
        for col, header in enumerate(headers):
            worksheet.write(0, col, header)
        for row, post in enumerate(posts, start=1):
            worksheet.write(row, 0, post.message_id)
            worksheet.write(row, 1, post.date)
            worksheet.write(row, 2, post.text)
            worksheet.write(row, 3, post.photo_url)
        workbook.close()

class DOCXOutputProcessor(OutputProcessor):
    def save(self, posts: List[TelegramPost], filename: str, delimiter: str = None):
        doc = Document()
        for post in posts:
            doc.add_heading(f"Message ID: {post.message_id}", level=2)
            doc.add_paragraph(f"Date: {post.date}")
            doc.add_paragraph(f"Text: {post.text}")
            if post.photo_url:
                doc.add_paragraph(f"Photo URL: {post.photo_url}")
            doc.add_paragraph("---")
        doc.save(filename)

class CSVOutputProcessor(OutputProcessor):
    def save(self, posts: List[TelegramPost], filename: str, delimiter: str = ","):
        with open(filename, "w", newline="", encoding="utf-8") as csvfile:
            writer = csv.writer(csvfile, delimiter=delimiter)
            writer.writerow(["Message ID", "Date", "Text", "Photo URL"])
            for post in posts:
                writer.writerow([post.message_id, post.date, post.text, post.photo_url])

class OutputProcessorFactory:
    @staticmethod
    def get_processor(output_format: str) -> OutputProcessor:
        processors = {
            "json": JSONOutputProcessor(),
            "txt": TXTOutputProcessor(),
            "markdown": MarkdownOutputProcessor(),
            "pdf": PDFOutputProcessor(),
            "html": HTMLOutputProcessor(),
            "xlsx": XLSXOutputProcessor(),
            "docx": DOCXOutputProcessor(),
            "csv": CSVOutputProcessor(),
        }
        return processors.get(output_format.lower(), JSONOutputProcessor())

DEFAULT_CONFIG = {
    "channels": [],
    "output_format": "json",
    "start_ids": {},
    "delimiter": "\n",
    "max_retries": 3,
    "retry_delay": 2,
    "timeout": 10,
    "max_empty_pages": 3
}

def create_default_config(config_file: str):
    with open(config_file, 'w') as f:
        yaml.dump(DEFAULT_CONFIG, f)
    logger.info(f"Created default configuration file: {config_file}")

def load_config(config_file: str) -> Dict[str, Any]:
    try:
        with open(config_file, 'r') as f:
            config = yaml.safe_load(f)
        return {**DEFAULT_CONFIG, **config}  # Merge with default config
    except FileNotFoundError:
        logger.warning(f"Configuration file not found: {config_file}")
        return DEFAULT_CONFIG
    except yaml.YAMLError as e:
        logger.error(f"Error parsing configuration file: {e}")
        return DEFAULT_CONFIG

def parse_arguments() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Telegram Channel Scraper")
    parser.add_argument("-c", "--config", type=str, default="config.yaml", help="Path to the configuration file")
    parser.add_argument("-f", "--format", type=str, choices=["txt", "json", "csv", "markdown", "pdf", "html", "xlsx", "docx"], help="Output format")
    parser.add_argument("-d", "--delimiter", type=str, help="Custom delimiter for TXT and CSV outputs")
    parser.add_argument("--channels", nargs='+', help="List of channel names to scrape")
    parser.add_argument("--create-config", action="store_true", help="Create a default configuration file")
    return parser.parse_args()

class TelegramChannelScraper:
    def __init__(self, channel_name: str, config: Dict[str, Any]):
        self.channel_name: str = channel_name
        self.base_url: str = f"https://t.me/s/{channel_name}"
        self.posts: Dict[str, TelegramPost] = {}
        self.oldest_id: Optional[str] = config['start_ids'].get(channel_name)
        self.newest_id: Optional[str] = None
        self.empty_page_count: int = 0
        self.max_empty_pages: int = config.get('max_empty_pages', 3)
        self.max_retries: int = config.get('max_retries', 3)
        self.retry_delay: int = config.get('retry_delay', 2)
        self.timeout: int = config.get('timeout', 10)

    async def get_page_content(self, session: aiohttp.ClientSession, before: Optional[str] = None) -> Optional[str]:
        url = self.base_url
        if before:
            url += f"?before={before}"
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
            "Cache-Control": "no-cache",
        }
        for attempt in range(self.max_retries):
            try:
                async with session.get(url, headers=headers, timeout=self.timeout) as response:
                    response.raise_for_status()
                    return await response.text()
            except (aiohttp.ClientError, asyncio.TimeoutError) as e:
                logger.warning(f"Request failed (attempt {attempt + 1}/{self.max_retries}): {e}")
                if attempt < self.max_retries - 1:
                    await asyncio.sleep(self.retry_delay)
                else:
                    logger.error(f"Failed to get page content after {self.max_retries} attempts")
        return None

    def parse_message(self, message_html: str) -> TelegramPost:
        parsers = [self.parse_with_bs4, self.parse_with_regex]
        result = TelegramPost(message_id="", date="", text="", photo_url="")
        for parser in parsers:
            parsed = parser(message_html)
            for key, value in parsed.items():
                if not getattr(result, key) and value:
                    setattr(result, key, value)
        return result

    def parse_with_bs4(self, message_html: str) -> Dict[str, str]:
        soup = BeautifulSoup(message_html, "html.parser")
        text_elem = soup.select_one(".tgme_widget_message_text")
        text = text_elem.get_text(strip=True) if text_elem else ""
        photo = soup.select_one(".tgme_widget_message_photo_wrap")
        photo_url = photo["style"].split("'")[1] if photo and "style" in photo.attrs else ""
        time_element = soup.select_one("time")
        date = time_element.get("datetime") if time_element else ""
        message_element = soup.select_one(".tgme_widget_message")
        message_id = message_element["data-post"].split("/")[-1] if message_element and "data-post" in message_element.attrs else ""
        return {"text": text, "photo_url": photo_url, "date": date, "message_id": message_id}

    def parse_with_regex(self, message_html: str) -> Dict[str, str]:
        text_pattern = re.compile(r'<div class="tgme_widget_message_text js-message_text" dir="auto">(.*?)</div>', re.DOTALL)
        photo_pattern = re.compile(r'<a class="tgme_widget_message_photo_wrap.*?background-image:url\(\'(.*?)\'\)')
        date_pattern = re.compile(r'<time datetime="(.*?)">')
        message_id_pattern = re.compile(r'data-post=".*?/(\d+)"')

        text = text_pattern.search(message_html)
        text = html.unescape(re.sub("<[^<]+?>", "", text.group(1))).strip() if text else ""

        photo = photo_pattern.search(message_html)
        photo_url = photo.group(1) if photo else ""

        date = date_pattern.search(message_html)
        date = date.group(1) if date else ""

        message_id = message_id_pattern.search(message_html)
        message_id = message_id.group(1) if message_id else ""

        return {"text": text, "photo_url": photo_url, "date": date, "message_id": message_id}

    async def scrape_channel(self, session: aiohttp.ClientSession, pbar: tqdm) -> None:
        while True:
            try:
                page_content = await self.get_page_content(session, self.oldest_id)
                if not page_content:
                    break

                soup = BeautifulSoup(page_content, "html.parser")
                messages = soup.select(".tgme_widget_message_wrap")

                if not messages:
                    self.empty_page_count += 1
                    if self.empty_page_count >= self.max_empty_pages:
                        logger.info(f"No new messages for {self.max_empty_pages} consecutive pages. Stopping.")
                        break
                    continue
                else:
                    self.empty_page_count = 0

                new_messages_count = 0
                for message in messages:
                    try:
                        parsed_message = self.parse_message(str(message))
                        message_id = parsed_message.message_id
                        if message_id and (parsed_message.text or parsed_message.photo_url):
                            if message_id not in self.posts:
                                self.posts[message_id] = parsed_message
                                new_messages_count += 1
                                pbar.update(1)
                                if not self.newest_id or int(message_id) > int(self.newest_id):
                                    self.newest_id = message_id
                                if not self.oldest_id or int(message_id) < int(self.oldest_id):
                                    self.oldest_id = message_id
                        else:
                            logger.info(f"Skipping empty or invalid message, ID: {message_id}")
                    except Exception as e:
                        logger.error(f"Error parsing message: {str(e)}")

                logger.info(f"Scraped {new_messages_count} new messages from {self.channel_name}. Total: {len(self.posts)}")

                if new_messages_count == 0:
                    self.empty_page_count += 1
                    if self.empty_page_count >= self.max_empty_pages:
                        logger.info(f"No new messages for {self.max_empty_pages} consecutive pages. Stopping.")
                        break
                else:
                    self.empty_page_count = 0

                if self.oldest_id and int(self.oldest_id) <= 1:
                    logger.info("Reached the earliest message. Stopping.")
                    break

                await asyncio.sleep(1)  # Be nice to the server
            except Exception as e:
                logger.error(f"Error during scraping process: {str(e)}")
                break

        logger.info(f"Scraping completed for {self.channel_name}. Total messages saved: {len(self.posts)}")

class TelegramMultiChannelScraper:
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.scrapers = [
            TelegramChannelScraper(channel, config)
            for channel in config['channels']
        ]
        self.output_processor = OutputProcessorFactory.get_processor(config['output_format'])

    async def scrape_all_channels(self) -> None:
        async with aiohttp.ClientSession() as session:
            total_messages = sum(len(scraper.posts) for scraper in self.scrapers)
            with tqdm(total=total_messages, desc="Scraping progress") as pbar:
                tasks = [scraper.scrape_channel(session, pbar) for scraper in self.scrapers]
                await asyncio.gather(*tasks)

    def save_results(self) -> None:
        all_posts = []
        for scraper in self.scrapers:
            all_posts.extend(scraper.posts.values())

        sorted_posts = sorted(all_posts, key=lambda x: x.date, reverse=True)

        filename = f"telegram_posts_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{self.config['output_format']}"
        try:
            self.output_processor.save(sorted_posts, filename, self.config.get('delimiter'))
            logger.info(f"Results saved to {filename}")
        except Exception as e:
            logger.error(f"Error saving results: {e}")

async def main():
    args = parse_arguments()

    if args.create_config:
        create_default_config(args.config)
        sys.exit(0)

    config = load_config(args.config)

    if args.format:
        config['output_format'] = args.format
    if args.delimiter:
        config['delimiter'] = args.delimiter
    if args.channels:
        config['channels'] = args.channels

    if not config['channels']:
        logger.error("No channels specified. Please provide channels in the config file or via command line.")
        sys.exit(1)

    multi_scraper = TelegramMultiChannelScraper(config)

    try:
        await multi_scraper.scrape_all_channels()
    except KeyboardInterrupt:
        logger.info("Program interrupted by user")
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
    finally:
        multi_scraper.save_results()
        logger.info("Program finished")

if __name__ == "__main__":
    asyncio.run(main())